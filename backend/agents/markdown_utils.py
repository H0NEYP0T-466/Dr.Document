"""Shared utilities for converting Markdown to LaTeX and python-docx."""
import re
from typing import List, Any


# ---------------------------------------------------------------------------
# LaTeX helpers
# ---------------------------------------------------------------------------

def latex_escape(text: str) -> str:
    """Escape special LaTeX characters (plain text only)."""
    _BS = '\x00BS\x00'
    text = text.replace('\\', _BS)
    text = text.replace('{', r'\{')
    text = text.replace('}', r'\}')
    text = text.replace(_BS, r'\textbackslash{}')
    text = text.replace('&', r'\&')
    text = text.replace('%', r'\%')
    text = text.replace('$', r'\$')
    text = text.replace('#', r'\#')
    text = text.replace('_', r'\_')
    text = text.replace('^', r'\textasciicircum{}')
    text = text.replace('~', r'\textasciitilde{}')
    text = text.replace('<', r'\textless{}')
    text = text.replace('>', r'\textgreater{}')
    return text


def inline_markdown_to_latex(text: str) -> str:
    """Apply inline markdown → LaTeX: bold, italic, inline code, then escape."""
    # We use placeholders to protect markdown spans from LaTeX escaping.
    saved: List[str] = []

    def _save(replacement: str) -> str:
        placeholder = f'\x00SPAN{len(saved)}\x00'
        saved.append(replacement)
        return placeholder

    # 1. Protect inline code first (highest priority)
    def _save_code(m: re.Match) -> str:
        return _save(r'\texttt{' + latex_escape(m.group(1)) + r'}')

    text = re.sub(r'`([^`]+)`', _save_code, text)

    # 2. Bold: **text** or __text__
    def _save_bold(m: re.Match) -> str:
        inner = m.group(1)
        # Recursively process inner content (may contain italic/code)
        return _save(r'\textbf{' + latex_escape(inner) + r'}')

    text = re.sub(r'\*\*([^*]+)\*\*', _save_bold, text)
    text = re.sub(r'__([^_]+)__', _save_bold, text)

    # 3. Italic: *text* (after bold so ** doesn't match here)
    #    For underscore-italic we only match _word_ with no surrounding underscores.
    def _save_italic(m: re.Match) -> str:
        inner = m.group(1)
        return _save(r'\textit{' + latex_escape(inner) + r'}')

    text = re.sub(r'\*([^*]+)\*', _save_italic, text)
    # Match _word_ only when surrounded by word boundaries or spaces
    text = re.sub(r'(?<!\w)_([^_\s][^_]*)_(?!\w)', _save_italic, text)

    # 4. Escape remaining plain text
    text = latex_escape(text)

    # 5. Restore all saved spans
    for idx, snippet in enumerate(saved):
        text = text.replace(f'\x00SPAN{idx}\x00', snippet)

    return text


def markdown_table_to_latex(table_lines: List[str]) -> str:
    """Convert a simple markdown table to a LaTeX tabular environment."""
    rows = []
    for line in table_lines:
        cells = [c.strip() for c in line.strip().strip('|').split('|')]
        rows.append(cells)

    if not rows:
        return ''

    col_count = len(rows[0])
    col_spec = '|' + 'l|' * col_count

    parts = [
        r'\begin{center}',
        r'\begin{tabular}{' + col_spec + r'}',
        r'\hline',
    ]
    for r_idx, row in enumerate(rows):
        escaped = [inline_markdown_to_latex(c) for c in row]
        parts.append(' & '.join(escaped) + r' \\')
        if r_idx == 0:
            parts.append(r'\hline')
    parts.append(r'\hline')
    parts.append(r'\end{tabular}')
    parts.append(r'\end{center}')
    return '\n'.join(parts)


def markdown_to_latex(text: str, subsection_cmd: str = 'subsection') -> str:
    """
    Convert Markdown constructs to LaTeX.

    Args:
        text: Markdown text to convert.
        subsection_cmd: LaTeX command to use for ## headings
                        ('subsection' for article, 'section' for report chapters).
    """
    lines = text.split('\n')
    output: List[str] = []
    i = 0
    while i < len(lines):
        line = lines[i]

        # Fenced code block
        if line.strip().startswith('```'):
            output.append(r'\begin{verbatim}')
            i += 1
            while i < len(lines) and not lines[i].strip().startswith('```'):
                output.append(lines[i])
                i += 1
            output.append(r'\end{verbatim}')
            i += 1
            continue

        # Markdown table
        if '|' in line and i + 1 < len(lines) and re.match(r'^[\s|:\-]+$', lines[i + 1]):
            table_lines = [line]
            i += 2  # skip separator
            while i < len(lines) and '|' in lines[i]:
                table_lines.append(lines[i])
                i += 1
            output.append(markdown_table_to_latex(table_lines))
            continue

        # Headings
        h3 = re.match(r'^###\s+(.*)', line)
        h2 = re.match(r'^##\s+(.*)', line)
        h1 = re.match(r'^#\s+(.*)', line)
        if h3:
            output.append(f'\\sub{subsection_cmd}{{{latex_escape(h3.group(1).strip())}}}')
            i += 1
            continue
        if h2:
            output.append(f'\\{subsection_cmd}{{{latex_escape(h2.group(1).strip())}}}')
            i += 1
            continue
        if h1:
            output.append(f'\\{subsection_cmd}{{{latex_escape(h1.group(1).strip())}}}')
            i += 1
            continue

        # Unordered list block
        if re.match(r'^[\-\*]\s+', line):
            output.append(r'\begin{itemize}')
            while i < len(lines) and re.match(r'^[\-\*]\s+', lines[i]):
                item_text = inline_markdown_to_latex(lines[i][2:].strip())
                output.append(f'  \\item {item_text}')
                i += 1
            output.append(r'\end{itemize}')
            continue

        # Ordered list block
        if re.match(r'^\d+\.\s+', line):
            output.append(r'\begin{enumerate}')
            while i < len(lines) and re.match(r'^\d+\.\s+', lines[i]):
                item_text = inline_markdown_to_latex(
                    re.sub(r'^\d+\.\s+', '', lines[i]).strip()
                )
                output.append(f'  \\item {item_text}')
                i += 1
            output.append(r'\end{enumerate}')
            continue

        # Normal text line
        output.append(inline_markdown_to_latex(line))
        i += 1

    return '\n'.join(output)


# ---------------------------------------------------------------------------
# python-docx helpers
# ---------------------------------------------------------------------------

def add_inline_markdown_runs(paragraph: Any, text: str) -> None:
    """Parse **bold**, *italic*, `code` and add styled runs to a paragraph."""
    from docx.shared import Pt

    pattern = re.compile(
        r'(\*\*[^*]+\*\*|__[^_]+__|'
        r'\*[^*]+\*|_[^_]+_|'
        r'`[^`]+`)'
    )
    segments = pattern.split(text)
    for seg in segments:
        if not seg:
            continue
        run = paragraph.add_run()
        if re.match(r'^\*\*(.+)\*\*$', seg) or re.match(r'^__(.+)__$', seg):
            run.text = seg[2:-2]
            run.bold = True
        elif re.match(r'^\*(.+)\*$', seg) or re.match(r'^_(.+)_$', seg):
            run.text = seg[1:-1]
            run.italic = True
        elif re.match(r'^`(.+)`$', seg):
            run.text = seg[1:-1]
            run.font.name = 'Courier New'
            run.font.size = Pt(10)
        else:
            run.text = seg


def add_markdown_table_to_docx(
    doc: Any, headers: List[str], rows: List[List[str]]
) -> None:
    """Add a markdown table as a Word table to the document."""
    from docx.shared import Pt

    if not headers:
        return
    col_count = len(headers)
    if rows:
        col_count = max(col_count, max(len(r) for r in rows))
    if col_count == 0:
        return

    table = doc.add_table(rows=1 + len(rows), cols=col_count)
    table.style = 'Table Grid'

    for j, cell_text in enumerate(headers):
        if j < col_count:
            cell = table.rows[0].cells[j]
            cell.text = ''
            run = cell.paragraphs[0].add_run(cell_text)
            run.bold = True
            run.font.name = 'Times New Roman'
            run.font.size = Pt(10)

    for r_idx, row in enumerate(rows):
        for j, cell_text in enumerate(row):
            if j < col_count:
                cell = table.rows[r_idx + 1].cells[j]
                cell.text = cell_text
                for run in cell.paragraphs[0].runs:
                    run.font.name = 'Times New Roman'
                    run.font.size = Pt(10)


def render_markdown_to_docx(doc: Any, content: str) -> None:
    """Render markdown content into a python-docx Document object."""
    from docx.shared import Pt

    lines = content.split('\n')
    i = 0
    while i < len(lines):
        line = lines[i]

        # Fenced code block
        if line.strip().startswith('```'):
            code_lines: List[str] = []
            i += 1
            while i < len(lines) and not lines[i].strip().startswith('```'):
                code_lines.append(lines[i])
                i += 1
            p = doc.add_paragraph('\n'.join(code_lines))
            for run in p.runs:
                run.font.name = 'Courier New'
                run.font.size = Pt(10)
            i += 1
            continue

        # Markdown table
        if '|' in line and i + 1 < len(lines) and re.match(r'^[\s|:\-]+$', lines[i + 1]):
            header_row = [c.strip() for c in line.strip().strip('|').split('|')]
            i += 2  # skip separator
            data_rows: List[List[str]] = []
            while i < len(lines) and '|' in lines[i]:
                data_rows.append(
                    [c.strip() for c in lines[i].strip().strip('|').split('|')]
                )
                i += 1
            add_markdown_table_to_docx(doc, header_row, data_rows)
            continue

        # Headings
        h3 = re.match(r'^###\s+(.*)', line)
        h2 = re.match(r'^##\s+(.*)', line)
        h1 = re.match(r'^#\s+(.*)', line)
        if h3:
            doc.add_heading(h3.group(1).strip(), level=4)
            i += 1
            continue
        if h2:
            doc.add_heading(h2.group(1).strip(), level=3)
            i += 1
            continue
        if h1:
            doc.add_heading(h1.group(1).strip(), level=2)
            i += 1
            continue

        # Unordered list
        if re.match(r'^[\-\*]\s+', line):
            while i < len(lines) and re.match(r'^[\-\*]\s+', lines[i]):
                item_text = lines[i][2:].strip()
                p = doc.add_paragraph(style='List Bullet')
                add_inline_markdown_runs(p, item_text)
                i += 1
            continue

        # Ordered list
        if re.match(r'^\d+\.\s+', line):
            while i < len(lines) and re.match(r'^\d+\.\s+', lines[i]):
                item_text = re.sub(r'^\d+\.\s+', '', lines[i]).strip()
                p = doc.add_paragraph(style='List Number')
                add_inline_markdown_runs(p, item_text)
                i += 1
            continue

        # Normal paragraph (skip blank lines)
        stripped = line.strip()
        if stripped:
            p = doc.add_paragraph()
            add_inline_markdown_runs(p, stripped)
            for run in p.runs:
                if not run.font.name or run.font.name == 'Courier New':
                    continue  # preserve code font
                run.font.name = 'Times New Roman'
                run.font.size = Pt(12)
        i += 1

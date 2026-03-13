"""Research Paper Formatter Agent - Assembles, polishes, generates LaTeX/PDF/DOCX"""
import asyncio
import os
import subprocess
import tempfile
from datetime import date
from typing import Dict, Any, List
from backend.agents.base_agent import BaseAgent
from backend.config import settings
from backend.logger import logger


class ResearchPaperFormatterAgent(BaseAgent):
    """
    Assembles all approved sections into a research paper.
    Produces: .tex, .pdf (via pdflatex), and .docx (via Node.js docx package).
    """

    def __init__(self):
        super().__init__("Research Paper Formatter Agent", settings.model_flash_lite)

    def process(self, input_data: Dict[str, Any]) -> Dict[str, Any]:
        """
        Format the research paper.

        Args:
            input_data: {
                'sections': List[{name, content}],
                'selected_title': str,
                'repo_name': str,
                'repo_url': str,
                'job_dir': str,       # absolute path to job directory
            }

        Returns:
            {
                'tex_content': str,
                'tex_path': str,
                'pdf_path': str | None,
                'docx_path': str | None,
                'latex_errors': str,
            }
        """
        sections = input_data.get('sections', [])
        selected_title = input_data.get('selected_title', 'Research Paper')
        repo_name = input_data.get('repo_name', 'Unknown')
        repo_url = input_data.get('repo_url', '')
        job_dir = input_data.get('job_dir', tempfile.gettempdir())

        logger.workflow_step("Research Paper Formatter", "Assembling paper")

        # Step 1: Polish content via LLM
        polished_sections = self._polish_content(sections, selected_title)

        # Step 2: Generate LaTeX
        tex_content = self._generate_latex(polished_sections, selected_title, repo_name, repo_url)
        tex_path = os.path.join(job_dir, 'research_paper.tex')
        with open(tex_path, 'w', encoding='utf-8') as f:
            f.write(tex_content)

        # Step 3: Compile PDF
        pdf_path, latex_errors = self._compile_pdf(tex_path, job_dir)

        # Log errors
        errors_log = os.path.join(job_dir, 'formatter_latex_errors.log')
        with open(errors_log, 'w', encoding='utf-8') as f:
            f.write(latex_errors or 'No errors')

        # Step 4: Generate DOCX
        docx_path = self._generate_docx(polished_sections, selected_title, repo_name, job_dir)

        return {
            'tex_content': tex_content,
            'tex_path': tex_path,
            'pdf_path': pdf_path,
            'docx_path': docx_path,
            'latex_errors': latex_errors,
        }

    def _polish_content(self, sections: List[Dict], title: str) -> List[Dict]:
        """Use LLM to fix grammar, consistency and tone across all sections.

        The system prompt explicitly forbids meta narration so the LLM output
        is always clean document text.  Any remaining chatter is stripped by
        the output validator as a safety net.
        """
        from backend.agents.output_validator import sanitize_content
        polished = []
        for sec in sections:
            messages = [
                {
                    "role": "system",
                    "content": (
                        "You are an academic copy-editor. Polish the provided section for "
                        "grammar, spelling, punctuation, consistent academic tone, and smooth "
                        "transitions. Do not change the substance or add new information. "
                        "IMPORTANT: Return ONLY the corrected section text — nothing else. "
                        "Do NOT say 'Here is the polished version', 'Certainly', 'Below is', "
                        "or include any preface, closing remark, or meta commentary. "
                        "Start the response with the first word of the section content."
                    ),
                },
                {
                    "role": "user",
                    "content": (
                        f"Paper: {title}\nSection: {sec['name']}\n\n{sec['content']}"
                    ),
                },
            ]
            try:
                polished_content = self._call_llm(messages, max_tokens=1500, temperature=0.2)
                polished_content = sanitize_content(polished_content.strip())
                polished.append({'name': sec['name'], 'content': polished_content})
            except Exception as e:
                logger.warning(f"Could not polish section '{sec['name']}': {e}")
                polished.append(sec)
        return polished

    def _generate_latex(
        self,
        sections: List[Dict],
        title: str,
        repo_name: str,
        repo_url: str,
    ) -> str:
        """Generate a complete IEEE-formatted LaTeX document."""
        # Escape title for LaTeX
        safe_title = self._latex_escape(title)
        author = repo_name.split('/')[0] if '/' in repo_name else repo_name

        # Build section bodies
        body_parts = []
        abstract_content = ''
        references_content = ''
        other_sections = []

        for sec in sections:
            name_lower = sec['name'].lower()
            if name_lower == 'abstract':
                abstract_content = self._latex_escape(sec['content'])
            elif name_lower == 'references':
                references_content = sec['content']
            else:
                other_sections.append(sec)

        for sec in other_sections:
            sec_title = self._latex_escape(sec['name'])
            sec_body = self._latex_escape(sec['content'])
            body_parts.append(f'\\section{{{sec_title}}}\n{sec_body}\n')

        body = '\n'.join(body_parts)

        # Build references
        ref_body = self._format_references_latex(references_content, repo_url)

        tex = f"""\\documentclass[conference]{{IEEEtran}}
\\usepackage{{cite}}
\\usepackage{{amsmath,amssymb,amsfonts}}
\\usepackage{{algorithmic}}
\\usepackage{{graphicx}}
\\usepackage{{textcomp}}
\\usepackage{{xcolor}}
\\usepackage{{hyperref}}
\\usepackage{{listings}}
\\usepackage{{booktabs}}

\\title{{{safe_title}}}
\\author{{{self._latex_escape(author)}}}
\\date{{\\today}}

\\begin{{document}}
\\maketitle

\\begin{{abstract}}
{abstract_content}
\\end{{abstract}}

{body}

\\begin{{thebibliography}}{{99}}
{ref_body}
\\end{{thebibliography}}

\\end{{document}}
"""
        return tex

    def _latex_escape(self, text: str) -> str:
        """Escape special LaTeX characters in plain text."""
        # Order matters — escape backslash first
        replacements = [
            ('\\', r'\textbackslash{}'),
            ('&', r'\&'),
            ('%', r'\%'),
            ('$', r'\$'),
            ('#', r'\#'),
            ('^', r'\textasciicircum{}'),
            ('~', r'\textasciitilde{}'),
        ]
        for old, new in replacements:
            text = text.replace(old, new)
        return text

    def _format_references_latex(self, references_content: str, repo_url: str) -> str:
        """Convert references content to LaTeX bibitem entries."""
        if not references_content.strip():
            return f'\\bibitem{{repo}} Repository: \\url{{{repo_url}}}'

        lines = []
        idx = 1
        for line in references_content.strip().splitlines():
            stripped = line.strip()
            if stripped and (stripped[0].isdigit() or stripped.startswith('[')):
                # Remove leading number/bracket
                clean = stripped.lstrip('0123456789[].) ')
                lines.append(f'\\bibitem{{ref{idx}}} {self._latex_escape(clean)}')
                idx += 1
        return '\n'.join(lines) if lines else f'\\bibitem{{repo}} Repository: \\url{{{repo_url}}}'

    def _compile_pdf(self, tex_path: str, job_dir: str) -> tuple:
        """Run pdflatex up to 3 times. Returns (pdf_path_or_None, error_log)."""
        pdf_path = tex_path.replace('.tex', '.pdf')
        errors = ''
        max_retries = 3

        for attempt in range(1, max_retries + 1):
            try:
                result = subprocess.run(
                    ['pdflatex', '-interaction=nonstopmode', '-output-directory', job_dir,
                     tex_path],
                    capture_output=True,
                    text=True,
                    timeout=120,
                )
                errors = result.stdout + result.stderr
                if result.returncode == 0 and os.path.exists(pdf_path):
                    logger.success(f"PDF compiled successfully on attempt {attempt}")
                    # Run second pass for references
                    subprocess.run(
                        ['pdflatex', '-interaction=nonstopmode', '-output-directory', job_dir,
                         tex_path],
                        capture_output=True,
                        text=True,
                        timeout=120,
                    )
                    return pdf_path, ''
                else:
                    logger.warning(f"pdflatex attempt {attempt} failed: {result.returncode}")
            except FileNotFoundError:
                logger.warning(
                    "pdflatex binary not found — install the system package, not the pip wrapper.\n"
                    "    Ubuntu/Debian : sudo apt-get install texlive-latex-base\n"
                    "    macOS         : brew install basictex\n"
                    "    Windows       : install MiKTeX from https://miktex.org"
                )
                return None, (
                    'pdflatex binary not found. '
                    'Install the system package (e.g. texlive-latex-base) — '
                    "'pip install pdflatex' only provides a Python wrapper and does not supply the binary."
                )
            except subprocess.TimeoutExpired:
                logger.warning(f"pdflatex timed out on attempt {attempt}")
                errors = 'pdflatex timed out'
            except Exception as e:
                logger.error(f"pdflatex error: {e}")
                errors = str(e)

        logger.warning(f"PDF compilation failed after {max_retries} attempts")
        return None, errors

    def _generate_docx(
        self,
        sections: List[Dict],
        title: str,
        repo_name: str,
        job_dir: str,
    ) -> str | None:
        """Generate DOCX using python-docx."""
        try:
            from docx import Document
            from docx.shared import Pt
            from docx.enum.text import WD_ALIGN_PARAGRAPH
        except ImportError:
            logger.warning("python-docx not installed — skipping DOCX generation")
            return None

        try:
            docx_path = os.path.join(job_dir, 'research_paper.docx')
            author = repo_name.split('/')[0] if '/' in repo_name else repo_name
            today = date.today().strftime('%B %d, %Y')

            doc = Document()

            # Set default font
            style = doc.styles['Normal']
            style.font.name = 'Times New Roman'
            style.font.size = Pt(12)

            # Cover page
            title_para = doc.add_paragraph()
            title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = title_para.add_run(title)
            run.bold = True
            run.font.size = Pt(24)
            run.font.name = 'Times New Roman'

            author_para = doc.add_paragraph()
            author_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = author_para.add_run(author)
            run.font.size = Pt(14)
            run.font.name = 'Times New Roman'

            date_para = doc.add_paragraph()
            date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = date_para.add_run(today)
            run.font.size = Pt(12)
            run.font.name = 'Times New Roman'

            gen_para = doc.add_paragraph()
            gen_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = gen_para.add_run('Generated by Dr. Document')
            run.italic = True
            run.font.size = Pt(10)
            run.font.name = 'Times New Roman'

            doc.add_page_break()

            # Sections
            for sec in sections:
                is_abstract = sec['name'].lower() == 'abstract'
                heading_level = 2 if is_abstract else 1
                doc.add_heading(sec['name'], level=heading_level)

                for paragraph in sec['content'].split('\n\n'):
                    cleaned = ' '.join(paragraph.split('\n')).strip()
                    if cleaned:
                        p = doc.add_paragraph()
                        run = p.add_run(cleaned)
                        run.font.name = 'Times New Roman'
                        run.font.size = Pt(12)

            doc.save(docx_path)
            logger.success("DOCX generated successfully")
            return docx_path
        except Exception as e:
            logger.error(f"DOCX generation error: {e}")
            return None

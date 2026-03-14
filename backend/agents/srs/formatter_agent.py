"""SRS Formatter Agent - Generates IEEE 830 SRS as LaTeX/PDF/DOCX"""
import os
import re
import subprocess
import tempfile
from datetime import date
from typing import Dict, Any, List
from backend.agents.base_agent import BaseAgent
from backend.agents.markdown_utils import (
    latex_escape,
    markdown_to_latex,
    render_markdown_to_docx,
)
from backend.config import settings
from backend.logger import logger


class SRSFormatterAgent(BaseAgent):
    """
    Assembles SRS sections into formatted documents.
    Produces: .tex, .pdf, .docx
    Includes revision history table in DOCX title page.
    """

    def __init__(self):
        super().__init__("SRS Formatter Agent", settings.model_flash_lite)

    def process(self, input_data: Dict[str, Any]) -> Dict[str, Any]:
        """
        Format the SRS document with correct IEEE 830 structure.

        Args:
            input_data: {
                'sections': List[{name, content}],
                'selected_title': str,
                'repo_name': str,
                'repo_url': str,
                'job_dir': str,
            }

        Returns:
            {
                'tex_path': str,
                'pdf_path': str | None,
                'docx_path': str | None,
                'latex_errors': str,
            }
        """
        sections = input_data.get('sections', [])
        selected_title = input_data.get('selected_title', 'Software Requirements Specification')
        repo_name = input_data.get('repo_name', 'Unknown')
        repo_url = input_data.get('repo_url', '')
        job_dir = input_data.get('job_dir', tempfile.gettempdir())

        logger.workflow_step("SRS Formatter", "Assembling SRS document")

        # Generate LaTeX
        tex_content = self._generate_latex(sections, selected_title, repo_name)
        tex_path = os.path.join(job_dir, 'srs.tex')
        with open(tex_path, 'w', encoding='utf-8') as f:
            f.write(tex_content)

        # Compile PDF
        pdf_path, latex_errors = self._compile_pdf(tex_path, job_dir)

        errors_log = os.path.join(job_dir, 'formatter_latex_errors.log')
        with open(errors_log, 'w', encoding='utf-8') as f:
            f.write(latex_errors or 'No errors')

        # Generate DOCX
        docx_path = self._generate_docx(sections, selected_title, repo_name, job_dir)

        return {
            'tex_path': tex_path,
            'pdf_path': pdf_path,
            'docx_path': docx_path,
            'latex_errors': latex_errors,
        }

    def _generate_latex(self, sections: List[Dict], title: str, repo_name: str) -> str:
        """Generate IEEE-style article LaTeX for SRS with markdown rendering."""
        safe_title = latex_escape(title)
        author = repo_name.split('/')[0] if '/' in repo_name else repo_name

        body_parts = []
        for sec in sections:
            if sec['name'].lower() in ('title page', 'table of contents'):
                continue
            sec_title = latex_escape(sec['name'])
            sec_body = markdown_to_latex(sec['content'], subsection_cmd='subsection')
            body_parts.append(f'\\section{{{sec_title}}}\n{sec_body}\n')

        body = '\n'.join(body_parts)

        tex = r"""\documentclass[12pt,a4paper]{article}
\usepackage{geometry}
\geometry{margin=1in}
\usepackage[utf8]{inputenc}
\usepackage[T1]{fontenc}
\usepackage{lmodern}
\usepackage{amsmath}
\usepackage{amssymb}
\usepackage{hyperref}
\usepackage{booktabs}
\usepackage{array}
\usepackage{longtable}
\usepackage{enumitem}
\usepackage{verbatim}
""" + f"""
\\title{{{safe_title}}}
\\author{{{latex_escape(author)}}}
\\date{{\\today}}

\\begin{{document}}
\\maketitle
\\thispagestyle{{empty}}

\\newpage
\\tableofcontents

\\newpage
{body}

\\end{{document}}
"""
        return tex

    def _compile_pdf(self, tex_path: str, job_dir: str) -> tuple:
        """Compile PDF with pdflatex (up to 3 attempts)."""
        pdf_path = tex_path.replace('.tex', '.pdf')
        errors = ''

        for attempt in range(1, 4):
            try:
                result = subprocess.run(
                    ['pdflatex', '-interaction=nonstopmode', '-output-directory', job_dir,
                     tex_path],
                    capture_output=True, text=True, timeout=120,
                )
                errors = result.stdout + result.stderr
                if result.returncode == 0 and os.path.exists(pdf_path):
                    logger.success(f"SRS PDF compiled on attempt {attempt}")
                    # Second pass for correct TOC page numbers
                    subprocess.run(
                        ['pdflatex', '-interaction=nonstopmode', '-output-directory', job_dir,
                         tex_path],
                        capture_output=True, text=True, timeout=120,
                    )
                    return pdf_path, ''
                else:
                    relevant = '\n'.join(
                        line for line in errors.splitlines()
                        if line.startswith('!') or 'Error' in line or 'error' in line
                    )
                    logger.warning(
                        f"pdflatex attempt {attempt} failed for SRS"
                        + (f':\n{relevant}' if relevant else f'\n{errors[-800:]}')
                    )
            except FileNotFoundError:
                logger.warning(
                    "pdflatex binary not found — install the system package, not the pip wrapper.\n"
                    "    Ubuntu/Debian : sudo apt-get install texlive-latex-base\n"
                    "    macOS         : brew install basictex\n"
                    "    Windows       : install MiKTeX from https://miktex.org"
                )
                return None, (
                    'pdflatex binary not found. '
                    'Install the system package (e.g. texlive-latex-base) — '
                    "'pip install pdflatex' only provides a Python wrapper and does not supply the binary."
                )
            except subprocess.TimeoutExpired:
                errors = 'pdflatex timed out'
            except Exception as e:
                errors = str(e)

        return None, errors

    def _generate_docx(
        self,
        sections: List[Dict],
        title: str,
        repo_name: str,
        job_dir: str,
    ) -> str | None:
        """Generate SRS DOCX with revision history table and markdown rendering."""
        try:
            from docx import Document
            from docx.shared import Pt
            from docx.enum.text import WD_ALIGN_PARAGRAPH
        except ImportError:
            logger.warning("python-docx not installed — skipping DOCX generation")
            return None

        try:
            docx_path = os.path.join(job_dir, 'srs.docx')
            author = repo_name.split('/')[0] if '/' in repo_name else repo_name
            today = date.today().strftime('%B %d, %Y')
            skip_sections = {'title page', 'table of contents'}

            doc = Document()

            # Default font
            style = doc.styles['Normal']
            style.font.name = 'Times New Roman'
            style.font.size = Pt(12)

            # Cover page
            title_para = doc.add_paragraph()
            title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = title_para.add_run(title)
            run.bold = True
            run.font.size = Pt(24)
            run.font.name = 'Times New Roman'

            author_para = doc.add_paragraph()
            author_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = author_para.add_run(author)
            run.font.size = Pt(14)
            run.font.name = 'Times New Roman'

            date_para = doc.add_paragraph()
            date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = date_para.add_run(today)
            run.font.size = Pt(12)
            run.font.name = 'Times New Roman'

            gen_para = doc.add_paragraph()
            gen_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = gen_para.add_run('Generated by Dr. Document')
            run.italic = True
            run.font.size = Pt(10)
            run.font.name = 'Times New Roman'

            # Revision history table
            doc.add_paragraph()
            rev_heading = doc.add_paragraph()
            run = rev_heading.add_run('Revision History')
            run.bold = True
            run.font.size = Pt(14)
            run.font.name = 'Times New Roman'

            table = doc.add_table(rows=2, cols=4)
            table.style = 'Table Grid'
            headers = ['Version', 'Date', 'Author', 'Description']
            for i, header in enumerate(headers):
                cell = table.rows[0].cells[i]
                cell.text = ''
                run = cell.paragraphs[0].add_run(header)
                run.bold = True
                run.font.name = 'Times New Roman'
                run.font.size = Pt(10)

            data_cells = table.rows[1].cells
            data_cells[0].text = '1.0'
            data_cells[1].text = today
            data_cells[2].text = author
            data_cells[3].text = 'Initial draft — auto-generated by Dr. Document'
            for cell in data_cells:
                for run in cell.paragraphs[0].runs:
                    run.font.name = 'Times New Roman'
                    run.font.size = Pt(10)

            doc.add_page_break()

            # Sections
            for sec in sections:
                if sec['name'].lower() in skip_sections:
                    continue
                is_subsection = (
                    bool(re.match(r'^\d+\.\d+', sec['name']))
                    or sec['name'].startswith('Appendix')
                )
                level = 2 if is_subsection else 1
                doc.add_heading(sec['name'], level=level)
                render_markdown_to_docx(doc, sec['content'])

            doc.save(docx_path)
            logger.success("SRS DOCX generated")
            return docx_path
        except Exception as e:
            logger.error(f"SRS DOCX error: {e}")
            return None
